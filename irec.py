import docx
import re
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional
import uuid

def count_words(text: str) -> int:
    """Counts words in a given text, ignoring whitespace and punctuation."""
    words = re.findall(r'\b\w+\b', text)
    return len(words)

def validate_date_format(date_str: str, format_str: str = "%m/%Y") -> Optional[datetime]:
    """Validates if date is in specified format and returns parsed datetime or None."""
    try:
        return datetime.strptime(date_str, format_str)
    except ValueError:
        return None

def validate_email(email: str) -> bool:
    """Validates if email contains '@' and '.'."""
    return bool(re.match(r".+@.+\..+", email))

def validate_file_name(file_name: str, pi_surname: str) -> Tuple[bool, str]:
    """
    Validates if a file name follows the Part 11 naming protocol.
    Expected formats:
    - Application: Surname_IREC Application_MMDDYYYY
    - Instruments: Surname_[Description-Language]_MMDDYYYY
    - Consent: Surname_[Description-Language]_MMDDYYYY
    - Recruitment: Surname_[Description-Language]_MMDDYYYY
    - Ethics Training: Surname_[TrainingBody_MMDDYYYY]
    - Other: Surname_[Description-Language]_MMDDYYYY
    """
    name_without_ext = file_name.rsplit('.', 1)[0]
    date_pattern = r"\d{2}\d{2}\d{4}"
    
    if re.match(rf"^{pi_surname}_IREC Application_{date_pattern}$", name_without_ext):
        return True, "Application form naming is valid."
    
    pattern = rf"^{pi_surname}_(.+?-(Eng|Ru|Kz))_{date_pattern}$"
    match = re.match(pattern, name_without_ext)
    if match:
        description = match.group(1)
        language = match.group(2)
        return True, f"File '{file_name}' naming is valid (Description: {description}, Language: {language})."
    
    training_pattern = rf"^{pi_surname}_(CITI|TRREE)_{date_pattern}$"
    if re.match(training_pattern, name_without_ext):
        return True, f"Ethics training certificate naming is valid for '{name_without_ext}'."
    
    return False, f"File '{file_name}' does not follow naming protocol."

def validate_part_0(doc: docx.Document) -> Tuple[Dict[str, List[str]], bool]:
    """Validates Part 0: Do I Submit an NU IREC Application?"""
    results = {"errors": [], "warnings": [], "info": []}
    part_0_text = ""
    in_part_0 = False
    
    for para in doc.paragraphs:
        if "Part 0: Do I Submit an NU IREC Application?" in para.text:
            in_part_0 = True
        if in_part_0:
            part_0_text += para.text + "\n"
        if "Part 1: Cover Sheet" in para.text:
            break
    
    if not part_0_text:
        results["errors"].append("Part 0 section not found.")
        return results, False
    
    questions = [
        {
            "text": "Does your research involve human subjects?",
            "pattern": r"Does your research involve human subjects.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: answer.startswith("Yes")
        },
        {
            "text": "Is this project being conducted solely to fulfill course requirements?",
            "pattern": r"Is this project being conducted solely to fulfill course requirements.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: answer.startswith("No")
        },
        {
            "text": "Is this project a quality assurance activity?",
            "pattern": r"Is this project a quality assurance activity.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: answer.startswith("No")
        },
        {
            "text": "Would you like to use this study to launch future investigations?",
            "pattern": r"Would you like to use this study to launch future investigations.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: answer.startswith("Yes")
        },
        {
            "text": "Would you like to disseminate or publish findings?",
            "pattern": r"Would you like to disseminate or publish findings.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: answer.startswith("Yes")
        },
        {
            "text": "Do you think this research is eligible for an Exemption?",
            "pattern": r"Do you think this research is eligible for an Exemption.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)",
            "application_needed": lambda answer: True
        }
    ]
    
    application_needed = False
    exemption_claimed = False
    for question in questions:
        match = re.search(question["pattern"], part_0_text, re.DOTALL)
        if not match:
            results["errors"].append(f"Response to '{question['text']}' not found or improperly formatted.")
            continue
        
        answer = match.group(1)
        if "☐" in answer:
            results["warnings"].append(f"Checkbox for '{question['text']}' is not marked (☐).")
            continue
        
        if question["application_needed"](answer):
            application_needed = True
        
        if question["text"] == "Do you think this research is eligible for an Exemption?" and answer == "Yes ☑":
            exemption_claimed = True
        
        results["info"].append(f"Response to '{question['text']}': {answer}")
    
    if not application_needed:
        results["errors"].append("Part 0 responses indicate no application is needed.")
    
    if exemption_claimed:
        justification = re.search(r"Outline the reasons why your study should be considered exempt:(.*?)(Part 1:|$)", part_0_text, re.DOTALL)
        if justification and justification.group(1).strip():
            results["info"].append("Exemption justification provided.")
        else:
            results["warnings"].append("Exemption claimed, but no justification provided.")
    
    exemption_categories = [
        ("f1", r"f1 (☑|☐) Research conducted in established or commonly accepted educational settings"),
        ("f2", r"f2 (☑|☐) Research involving the use of educational tests"),
        ("f3", r"f3 (☑|☐) Research involving the collection or study of existing data")
    ]
    
    for category_name, pattern in exemption_categories:
        match = re.search(pattern, part_0_text, re.DOTALL)
        if not match:
            results["errors"].append(f"Exemption category '{category_name}' not found.")
            continue
        if match.group(1) == "☑":
            results["errors"].append(f"Exemption category '{category_name}' is checked. Must be unchecked (☐).")
        else:
            results["info"].append(f"Exemption category '{category_name}' is correctly unchecked (☐).")
    
    return results, exemption_claimed

def validate_part_1(doc: docx.Document, exemption_claimed: bool) -> Dict[str, List[str]]:
    """Validates Part 1: Cover Sheet."""
    results = {"errors": [], "warnings": [], "info": []}
    part_1_text = ""
    in_part_1 = False
    
    for para in doc.paragraphs:
        if "Part 1: Cover Sheet" in para.text:
            in_part_1 = True
        if in_part_1:
            part_1_text += para.text + "\n"
        if "Part 2: Research Team Details" in para.text:
            break
    
    if not part_1_text:
        results["errors"].append("Part 1: Cover Sheet section not found.")
        return results
    
    fields = [
        ("Principal Investigator:", r"Principal Investigator:\s*([^\n]+)"),
        ("Application Date:", r"Application Date:\s*([^\n]+)"),
        ("Nazarbayev University Unit \(School\):", r"Nazarbayev University Unit \(School\):\s*([^\n]+)"),
        ("Primary Research Discipline:", r"Primary Research Discipline:\s*([^\n]+)"),
        ("Application Title:", r"Application Title:\s*([^\n]+)")
    ]
    
    for field_name, pattern in fields:
        match = re.search(pattern, part_1_text)
        if not match or not match.group(1).strip():
            results["errors"].append(f"Field '{field_name}' is missing or empty.")
        else:
            value = match.group(1).strip()
            results["info"].append(f"Field '{field_name}' filled: {value}")
            if field_name == "Application Date:":
                try:
                    datetime.strptime(value, "%m/%d/%Y")
                    results["info"].append("Application Date is in valid format (MM/DD/YYYY).")
                except ValueError:
                    results["errors"].append("Application Date is not in valid format (MM/DD/YYYY).")
    
    review_types = [
        ("An Expedited Review", r"An Expedited Review\s+.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("A Full Board Review", r"A Full Board Review\s+.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("An Exemption", r"An Exemption\s+.*?(Yes ☑|No ☑|Yes ☐|No ☐)")
    ]
    
    selected_reviews = []
    for review_name, pattern in review_types:
        match = re.search(pattern, part_1_text, re.DOTALL)
        if not match:
            results["errors"].append(f"Response to '{review_name}' not found or improperly formatted.")
            continue
        
        answer = match.group(1)
        if "☐" in answer:
            results["warnings"].append(f"Checkbox for '{review_name}' is not marked (☐).")
        elif answer == "Yes ☑":
            selected_reviews.append(review_name)
        results["info"].append(f"Response to '{review_name}': {answer}")
    
    if len(selected_reviews) != 1:
        results["errors"].append(f"Exactly one review type must be selected. Found {len(selected_reviews)}.")
    elif selected_reviews[0] != "An Expedited Review":
        results["errors"].append(f"School-level review requires 'An Expedited Review'. Selected: {selected_reviews[0]}.")
    
    if "An Exemption" in selected_reviews and not exemption_claimed:
        results["errors"].append("Exemption selected in Part 1, but Part 0 does not claim exemption.")
    
    return results

def validate_part_2(doc: docx.Document) -> Tuple[Dict[str, List[str]], str]:
    """Validates Part 2: Research Team Details and extracts PI surname."""
    results = {"errors": [], "warnings": [], "info": []}
    part_2_text = ""
    in_part_2 = False
    pi_surname = ""
    
    for para in doc.paragraphs:
        if "Part 2: Research Team Details" in para.text:
            in_part_2 = True
        if in_part_2:
            part_2_text += para.text + "\n"
        if "Part 3: Research Design" in para.text:
            break
    
    if not part_2_text:
        results["errors"].append("Part 2: Research Team Details section not found.")
        return results, pi_surname
    
    pi_name_match = re.search(r"Principal Investigator\s*\n\s*Name:\s*([^\n]+)", part_2_text, re.DOTALL)
    if pi_name_match and pi_name_match.group(1).strip():
        pi_name = pi_name_match.group(1).strip()
        pi_surname = pi_name.split()[-1]  # Assume last word is surname
        results["info"].append(f"PI Name: {pi_name}, Surname extracted: {pi_surname}")
    
    pi_fields = [
        ("PI Name:", r"Principal Investigator\s*\n\s*Name:\s*([^\n]+)"),
        ("PI NU ID:", r"Principal Investigator\s*\n.*?\n\s*NU ID:\s*([^\n]+)"),
        ("PI NU School:", r"Principal Investigator\s*\n.*?\n\s*NU School:\s*([^\n]+)"),
        ("PI Department:", r"Principal Investigator\s*\n.*?\n\s*Department:\s*([^\n]+)"),
        ("PI Position:", r"Principal Investigator\s*\n.*?\n\s*Position:\s*([^\n]+)"),
        ("PI E-mail address:", r"Principal Investigator\s*\n.*?\n\s*E-mail address:\s*([^\n]+)"),
        ("PI Daytime Phone:", r"Principal Investigator\s*\n.*?\n\s*Daytime Phone:\s*([^\n]+)"),
        ("PI Mobile phone:", r"Principal Investigator\s*\n.*?\n\s*Mobile phone:\s*([^\n]+)"),
        ("PI CITI Training completion date:", r"Principal Investigator\s*\n.*?\n\s*CITI Training completion date:\s*([^\n]+)")
    ]
    
    for field_name, pattern in pi_fields:
        match = re.search(pattern, part_2_text, re.DOTALL)
        if not match or not match.group(1).strip():
            results["errors"].append(f"Field '{field_name}' is missing or empty.")
        else:
            value = match.group(1).strip()
            results["info"].append(f"Field '{field_name}' filled: {value}")
            if field_name == "PI E-mail address:" and not validate_email(value):
                results["errors"].append(f"Invalid email format for '{field_name}': {value}")
            if field_name == "PI CITI Training completion date:":
                try:
                    citi_date = datetime.strptime(value, "%m/%d/%Y")
                    three_years_ago = datetime.now() - timedelta(days=3*365)
                    if citi_date < three_years_ago:
                        results["errors"].append("PI CITI Training date is older than 3 years.")
                    else:
                        results["info"].append("PI CITI Training date is valid.")
                except ValueError:
                    results["errors"].append("PI CITI Training date is not in valid format (MM/DD/YYYY).")
    
    pi_citi_status = re.search(r"Principal Investigator\s*\n.*?\n\s*Have you completed the CITI basic course.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_2_text, re.DOTALL)
    if not pi_citi_status:
        results["errors"].append("PI CITI training status not found.")
    elif pi_citi_status.group(1) == "No ☑":
        results["errors"].append("PI CITI training status is 'No'.")
    elif pi_citi_status.group(1) in ["Yes ☐", "No ☐"]:
        results["warnings"].append("PI CITI training status checkbox is not marked (☐).")
    else:
        results["info"].append("PI CITI training status is 'Yes'.")
    
    ra_fields = [
        ("RA Name:", r"Research Advisor:\s*\n\s*Name:\s*([^\n]+)"),
        ("RA NU ID:", r"Research Advisor:\s*\n.*?\n\s*NU ID:\s*([^\n]+)"),
        ("RA NU School:", r"Research Advisor:\s*\n.*?\n\s*NU School:\s*([^\n]+)"),
        ("RA Department:", r"Research Advisor:\s*\n.*?\n\s*Department:\s*([^\n]+)"),
        ("RA Position:", r"Research Advisor:\s*\n.*?\n\s*Position:\s*([^\n]+)"),
        ("RA E-mail address:", r"Research Advisor:\s*\n.*?\n\s*E-mail address:\s*([^\n]+)"),
        ("RA CITI or alternative training completion date:", r"Research Advisor:\s*\n.*?\n\s*CITI or alternative training completion date:\s*([^\n]+)")
    ]
    
    for field_name, pattern in ra_fields:
        match = re.search(pattern, part_2_text, re.DOTALL)
        if not match or not match.group(1).strip():
            results["errors"].append(f"Field '{field_name}' is missing or empty.")
        else:
            value = match.group(1).strip()
            results["info"].append(f"Field '{field_name}' filled: {value}")
            if field_name == "RA E-mail address:" and not validate_email(value):
                results["errors"].append(f"Invalid email format for '{field_name}': {value}")
            if field_name == "RA CITI or alternative training completion date:":
                try:
                    citi_date = datetime.strptime(value, "%m/%d/%Y")
                    three_years_ago = datetime.now() - timedelta(days=3*365)
                    if citi_date < three_years_ago:
                        results["errors"].append("RA CITI training date is older than 3 years.")
                    else:
                        results["info"].append("RA CITI training date is valid.")
                except ValueError:
                    results["errors"].append("RA CITI training date is not in valid format (MM/DD/YYYY).")
    
    ra_citi_status = re.search(r"Research Advisor:\s*\n.*?\n\s*Have you completed the CITI basic course.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_2_text, re.DOTALL)
    if not ra_citi_status:
        results["errors"].append("RA CITI training status not found.")
    elif ra_citi_status.group(1) == "No ☑":
        results["errors"].append("RA CITI training status is 'No'.")
    elif ra_citi_status.group(1) in ["Yes ☐", "No ☐"]:
        results["warnings"].append("RA CITI training status checkbox is not marked (☐).")
    else:
        results["info"].append("RA CITI training status is 'Yes'.")
    
    additional_investigator_pattern = r"Additional Investigator\(s\):.*?\n\s*Name:\s*([^\n]*)\n\s*NU ID:\s*([^\n]*)\n\s*NU School:\s*([^\n]*)\n\s*Department:\s*([^\n]*)\n\s*Position:\s*([^\n]*)\n\s*E-mail address:\s*([^\n]*)\n\s*Have you completed the CITI basic course.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)\n\s*.*?\n\s*CITI or alternative training completion date:\s*([^\n]*)"
    additional_investigators = re.finditer(additional_investigator_pattern, part_2_text, re.DOTALL)
    
    investigator_count = 0
    for match in additional_investigators:
        investigator_count += 1
        name = match.group(1).strip()
        if name:
            fields = [
                ("AI Name", name),
                ("AI NU ID", match.group(2).strip()),
                ("AI NU School", match.group(3).strip()),
                ("AI Department", match.group(4).strip()),
                ("AI Position", match.group(5).strip()),
                ("AI E-mail address", match.group(6).strip()),
                ("AI CITI or alternative training completion date", match.group(8).strip())
            ]
            
            for field_name, value in fields:
                if not value:
                    results["errors"].append(f"Additional Investigator {investigator_count}: Field '{field_name}' is missing or empty.")
                else:
                    results["info"].append(f"Additional Investigator {investigator_count}: Field '{field_name}' filled: {value}")
                    if field_name == "AI E-mail address:" and not validate_email(value):
                        results["errors"].append(f"Additional Investigator {investigator_count}: Invalid email format for '{field_name}': {value}")
                    if field_name == "AI CITI or alternative training completion date":
                        try:
                            citi_date = datetime.strptime(value, "%m/%d/%Y")
                            three_years_ago = datetime.now() - timedelta(days=3*365)
                            if citi_date < three_years_ago:
                                results["errors"].append(f"Additional Investigator {investigator_count}: CITI training date is older than 3 years.")
                            else:
                                results["info"].append(f"Additional Investigator {investigator_count}: CITI training date is valid.")
                        except ValueError:
                            results["errors"].append(f"Additional Investigator {investigator_count}: CITI training date is not in valid format (MM/DD/YYYY).")
            
            citi_status = match.group(7)
            if citi_status == "No ☑":
                results["errors"].append(f"Additional Investigator {investigator_count}: CITI training status is 'No'.")
            elif citi_status in ["Yes ☐", "No ☐"]:
                results["warnings"].append(f"Additional Investigator {investigator_count}: CITI training status checkbox is not marked (☐).")
            else:
                results["info"].append(f"Additional Investigator {investigator_count}: CITI training status is 'Yes'.")
    
    if investigator_count == 0:
        results["info"].append("No Additional Investigators specified.")
    
    student_section = re.search(r"For students:\s*\n\s*Undergraduate (☑|☐)\s*Masters (☑|☐)\s*PhD (☑|☐)\s*Other (☑|☐)\s*\n\s*Course:\s*([^\n]*)", part_2_text, re.DOTALL)
    if not student_section:
        results["errors"].append("For students section not found or improperly formatted.")
    else:
        undergraduate, masters, phd, other, course = student_section.groups()
        selected_categories = [cat for cat, checked in [
            ("Undergraduate", undergraduate),
            ("Masters", masters),
            ("PhD", phd),
            ("Other", other)
        ] if checked == "☑"]
        
        if len(selected_categories) != 1:
            results["errors"].append(f"Exactly one student category must be selected. Found {len(selected_categories)}.")
        else:
            results["info"].append(f"Student category selected: {selected_categories[0]}.")
        
        if not course.strip():
            results["errors"].append("Course field in For students section is missing or empty.")
        else:
            results["info"].append(f"Course field filled: {course.strip()}.")
    
    return results, pi_surname

def validate_part_3(doc: docx.Document) -> Tuple[Dict[str, List[str]], List[Dict], str]:
    """Validates Part 3: Research Design and collects methodology text."""
    results = {"errors": [], "warnings": [], "info": []}
    required_forms = [
        {"form": "Appendix A: IREC Application Form", "reason": "Required for all submissions."},
        {"form": "CITI Training Certificates", "reason": "Required for all team members."}
    ]
    part_3_text = ""
    in_part_3 = False
    
    for para in doc.paragraphs:
        if "Part 3: Research Design" in para.text:
            in_part_3 = True
        if in_part_3:
            part_3_text += para.text + "\n"
        if "Part 4: Participants" in para.text:
            break
    
    if not part_3_text:
        results["errors"].append("Part 3: Research Design section not found.")
        return results, required_forms, ""
    
    fields = [
        ("Purpose of the research", r"What is the purpose of the research\?.*?\n(.*?)(What question\(s\) do you hope to answer\?|$)", 250, 300),
        ("Research question(s)", r"What question\(s\) do you hope to answer\?.*?\n(.*?)(Describe the data collection methodology|$)", None, None),
        ("Data collection methodology", r"Describe the data collection methodology.*?\n(.*?)(Briefly describe the data analysis processes|$)", 250, 300),
        ("Data analysis processes", r"Briefly describe the data analysis processes.*?\n(.*?)(Briefly describe the research sites|$)", 150, 300),
        ("Research sites", r"Briefly describe the research sites.*?\n(.*?)(Part 4:|$)", None, None)
    ]
    
    methodology_text = ""
    for field_name, pattern, min_words, max_words in fields:
        match = re.search(pattern, part_3_text, re.DOTALL)
        if not match or not match.group(1).strip():
            results["errors"].append(f"Field '{field_name}' is missing or empty.")
        else:
            value = match.group(1).strip()
            results["info"].append(f"Field '{field_name}' filled.")
            if field_name == "Data collection methodology":
                methodology_text = value
            if min_words and max_words:
                word_count = count_words(value)
                if word_count < min_words or word_count > max_words:
                    results["warnings"].append(f"Field '{field_name}' has {word_count} words, expected {min_words}–{max_words}.")
                else:
                    results["info"].append(f"Field '{field_name}' word count is valid: {word_count}.")
    
    methodology = methodology_text.lower()
    research_sites = re.search(r"Briefly describe the research sites.*?\n(.*?)(Part 4:|$)", part_3_text, re.DOTALL)
    research_sites_text = research_sites.group(1).lower() if research_sites and research_sites.group(1).strip() else ""
    languages = ["English"]
    if "kazakhstan" in research_sites_text:
        languages.extend(["Russian", "Kazakh"])
    elif research_sites_text and "nazarbayev university" not in research_sites_text:
        languages.append("Official language(s) of the country")
    
    if any(term in methodology for term in ["interview", "focus group", "observation", "action research"]):
        required_forms.extend([
            {"form": "Appendix B: Written Informed Consent Form", "reason": f"Required for qualitative research in {', '.join(languages)}."},
            {"form": "Appendix D: Oral Consent Script", "reason": f"Required if oral consent is used in {', '.join(languages)}."},
            {"form": "Interview Questions/Focus Group Guides", "reason": "Required for qualitative methods."},
            {"form": "Recruitment Materials (e.g., emails, flyers)", "reason": "Required for participant notification."}
        ])
    
    if any(term in methodology for term in ["survey", "clinical trial", "existing data set", "human genetics"]):
        if "internet survey" in methodology or "online survey" in methodology:
            required_forms.append({"form": "Appendix C: Informed Consent Form for Internet Surveys", "reason": f"Required for internet surveys in {', '.join(languages)}."})
        else:
            required_forms.append({"form": "Appendix B: Written Informed Consent Form", "reason": f"Required for quantitative research in {', '.join(languages)}."})
        required_forms.append({"form": "Surveys/Questionnaires", "reason": "Required for quantitative methods."})
    
    if "mixed method" in methodology:
        required_forms.extend([
            {"form": "Appendix B: Written Informed Consent Form", "reason": f"Required for mixed methods in {', '.join(languages)}."},
            {"form": "Recruitment Materials (e.g., emails, flyers)", "reason": "Required for participant notification."}
        ])
        if "interview" in methodology or "focus group" in methodology:
            required_forms.extend([
                {"form": "Appendix D: Oral Consent Script", "reason": f"Required if oral consent is used in {', '.join(languages)}."},
                {"form": "Interview Questions/Focus Group Guides", "reason": "Required for qualitative components."}
            ])
        if "survey" in methodology:
            required_forms.append({"form": "Surveys/Questionnaires", "reason": "Required for quantitative components."})
    
    if "genetic" in methodology or "biobank" in methodology:
        required_forms.append({"form": "Appendix M: Written Informed Consent Form For Genetic and/or Biobank Research", "reason": f"Required for genetic/biobank research in {', '.join(languages)}."})
    
    if "collaborator" in methodology or "external organization" in methodology:
        required_forms.append({"form": "Appendix L: Confidentiality Agreement Form", "reason": f"Required for external collaborators in {', '.join(languages)}."})
    
    if research_sites_text and "nazarbayev university" not in research_sites_text:
        required_forms.append({"form": "Letters of Support/Approval from Outside Organizations", "reason": "Required for external sites."})
    
    if "visual stimuli" in methodology:
        required_forms.append({"form": "Visual Stimuli", "reason": "Required if visual stimuli are presented."})
    
    if "attach" in part_3_text.lower() or "appendix" in part_3_text.lower():
        results["info"].append("References to attachments detected in Part 3.")
    else:
        results["warnings"].append("No references to attachments detected in Part 3.")
    
    return results, required_forms, methodology_text

def validate_part_4(doc: docx.Document, required_forms: List[Dict]) -> Tuple[Dict[str, List[str]], List[Dict]]:
    """Validates Part 4: Participants."""
    results = {"errors": [], "warnings": [], "info": []}
    part_4_text = ""
    in_part_4 = False
    
    for para in doc.paragraphs:
        if "Part 4: Participants" in para.text:
            in_part_4 = True
        if in_part_4:
            part_4_text += para.text + "\n"
        if "Part 5: Detailed Procedures" in para.text:
            break
    
    if not part_4_text:
        results["errors"].append("Part 4: Participants section not found.")
        return results, required_forms
    
    special_populations = [
        ("Minors", r"Minors \(under 18 years of age\)\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Legally incompetent", r"Legally incompetent\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Prisoners", r"Prisoners\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Perinatal women", r"Perinatal women.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Institutionalized", r"Institutionalized\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Mentally incapacitated", r"Mentally incapacitated\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Sexual behaviors", r"Sexual behaviors\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Drug use", r"Drug use\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Illegal conduct", r"Illegal conduct\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
        ("Use of alcohol", r"Use of alcohol\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)")
    ]
    
    special_populations_yes = []
    for pop_name, pattern in special_populations:
        match = re.search(pattern, part_4_text, re.DOTALL)
        if not match:
            results["errors"].append(f"Response to '{pop_name}' not found or improperly formatted.")
            continue
        
        answer = match.group(1)
        if "☐" in answer:
            results["warnings"].append(f"Checkbox for '{pop_name}' is not marked (☐).")
        elif answer == "Yes ☑":
            special_populations_yes.append(pop_name)
            results["info"].append(f"Special population '{pop_name}' selected: Yes.")
            if pop_name == "Minors":
                required_forms.extend([
                    {"form": "Appendix E: Assent Form", "reason": "Required for minors in English, Russian, Kazakh."},
                    {"form": "Parental Consent Forms", "reason": "Required for minors in English, Russian, Kazakh."}
                ])
            else:
                required_forms.append({"form": "Appendix B: Written Informed Consent Form", "reason": f"Required for special population '{pop_name}' in English, Russian, Kazakh."})
                if pop_name in ["Sexual behaviors", "Drug use", "Illegal conduct", "Use of alcohol"]:
                    required_forms.append({"form": "Appendix L: Confidentiality Agreement Form", "reason": f"Recommended for sensitive subjects ('{pop_name}')."})
        else:
            results["info"].append(f"Special population '{pop_name}' selected: No.")
    
    other_special = re.search(r"Other \(please specify\)\s*([^\n]*)", part_4_text)
    if other_special and other_special.group(1).strip():
        results["info"].append(f"Other special population: {other_special.group(1).strip()}.")
        required_forms.append({"form": "Appendix B: Written Informed Consent Form", "reason": "Required for other special populations."})
    
    sample_size = re.search(r"Expected number of participants or sample size:\s*(\d+)", part_4_text)
    if not sample_size or not sample_size.group(1):
        results["errors"].append("Sample size is missing or invalid.")
    else:
        results["info"].append(f"Sample size: {sample_size.group(1)}.")
    
    participant_fields = [
        ("Languages of communication", r"Languages of communication:\s*([^\n]+)"),
        ("Gender, race or ethnic group", r"Gender, race or ethnic group.*?:\s*([^\n]+)"),
        ("Affiliation of participants", r"Affiliation of participants.*?:\s*([^\n]+)"),
        ("Mental health", r"Participants’ general state of mental health:\s*([^\n]+)"),
        ("Physical health", r"Participants’ general state of physical health:\s*([^\n]+)")
    ]
    
    for field_name, pattern in participant_fields:
        match = re.search(pattern, part_4_text)
        if not match or not match.group(1).strip():
            results["errors"].append(f"Field '{field_name}' is missing or empty.")
        else:
            results["info"].append(f"Field '{field_name}' filled: {match.group(1).strip()}.")
    
    justification_na = re.search(r"Explain why you have chosen this particular group.*?\n.*?(N/A ☑|N/A ☐)", part_4_text, re.DOTALL)
    justification_text = re.search(r"Explain why you have chosen this particular group.*?\n(.*?)(What is your relationship to the participants\?|$)", part_4_text, re.DOTALL)
    
    if justification_na and justification_na.group(1) == "N/A ☑":
        if special_populations_yes:
            results["errors"].append("Justification cannot be N/A when special populations are selected.")
        else:
            results["info"].append("Justification marked as N/A.")
    elif justification_text and justification_text.group(1).strip():
        results["info"].append("Justification provided.")
    else:
        results["errors"].append("Justification is missing or empty.")
    
    relationship = re.search(r"What is your relationship to the participants\?.*?\n(.*?)(Does your relationship potentially create any power|$)", part_4_text, re.DOTALL)
    power_dynamics = re.search(r"Does your relationship potentially create any power.*?\n(.*?)(\n|$)", part_4_text, re.DOTALL)
    
    if not relationship or not relationship.group(1).strip():
        results["errors"].append("Relationship to participants is missing or empty.")
    else:
        results["info"].append(f"Relationship: {relationship.group(1).strip()}.")
    
    if not power_dynamics or not power_dynamics.group(1).strip():
        results["errors"].append("Power dynamics description is missing or empty.")
    else:
        results["info"].append(f"Power dynamics: {power_dynamics.group(1).strip()}.")
    
    recruitment = re.search(r"Will participants be recruited\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_4_text)
    contact_method_na = re.search(r"How will you contact potential participants.*?\n.*?(N/A ☑|N/A ☐)", part_4_text, re.DOTALL)
    contact_method_text = re.search(r"How will you contact potential participants.*?\n(.*?)(Describe the method for recruiting participants|$)", part_4_text, re.DOTALL)
    recruitment_method_na = re.search(r"Describe the method for recruiting participants.*?\n.*?(N/A ☑|N/A ☐)", part_4_text, re.DOTALL)
    recruitment_method_text = re.search(r"Describe the method for recruiting participants.*?\n(.*?)(Exclusions:|$)", part_4_text, re.DOTALL)
    
    if not recruitment:
        results["errors"].append("Recruitment question not found or improperly formatted.")
    else:
        recruitment_answer = recruitment.group(1)
        if "☐" in recruitment_answer:
            results["warnings"].append("Recruitment checkbox is not marked (☐).")
        elif recruitment_answer == "Yes ☑":
            results["info"].append("Participants will be recruited: Yes.")
            required_forms.append({"form": "Recruitment Materials (e.g., emails, flyers)", "reason": "Required for recruitment."})
            if contact_method_na and contact_method_na.group(1) == "N/A ☑":
                results["errors"].append("Contact method cannot be N/A when recruitment is Yes.")
            elif not contact_method_text or not contact_method_text.group(1).strip():
                results["errors"].append("Contact method description is missing or empty.")
            else:
                results["info"].append("Contact method description provided.")
            
            if recruitment_method_na and recruitment_method_na.group(1) == "N/A ☑":
                results["errors"].append("Recruitment method cannot be N/A when recruitment is Yes.")
            elif not recruitment_method_text or not recruitment_method_text.group(1).strip():
                results["errors"].append("Recruitment method description is missing or empty.")
            else:
                results["info"].append("Recruitment method description provided.")
        else:
            results["info"].append("Participants will be recruited: No.")
            if (contact_method_na and contact_method_na.group(1) != "N/A ☑") or (contact_method_text and contact_method_text.group(1).strip()):
                results["errors"].append("Contact method should be N/A when recruitment is No.")
            if (recruitment_method_na and recruitment_method_na.group(1) != "N/A ☑") or (recruitment_method_text and recruitment_method_text.group(1).strip()):
                results["errors"].append("Recruitment method should be N/A when recruitment is No.")
    
    exclusions_na = re.search(r"Exclusions:.*?\n.*?(N/A ☑|N/A ☐)", part_4_text, re.DOTALL)
    exclusions_text = re.search(r"Exclusions:.*?\n(.*?)(Procedures in the event of a participant withdrawing|$)", part_4_text, re.DOTALL)
    
    if exclusions_na and exclusions_na.group(1) == "N/A ☑":
        results["info"].append("Exclusions marked as N/A.")
    elif exclusions_text and exclusions_text.group(1).strip():
        results["info"].append("Exclusions description provided.")
    else:
        results["errors"].append("Exclusions description is missing or empty.")
    
    withdrawal = re.search(r"Procedures in the event of a participant withdrawing.*?\n(.*?)(Part 5:|$)", part_4_text, re.DOTALL)
    if not withdrawal or not withdrawal.group(1).strip():
        results["errors"].append("Withdrawal procedures description is missing or empty.")
    else:
        results["info"].append("Withdrawal procedures description provided.")
    
    return results, required_forms

def validate_part_5(doc: docx.Document, required_forms: List[Dict]) -> Tuple[Dict[str, List[str]], List[Dict], str]:
    """Validates Part 5: Detailed Procedures."""
    results = {"errors": [], "warnings": [], "info": []}
    part_5_text = ""
    in_part_5 = False
    
    for para in doc.paragraphs:
        if "Part 5: Detailed Procedures" in para.text:
            in_part_5 = True
        if in_part_5:
            part_5_text += para.text + "\n"
        if "Part 6: Data Management Plan" in para.text:
            break
    
    if not part_5_text:
        results["errors"].append("Part 5: Detailed Procedures section not found.")
        return results, required_forms, ""
    
    dates = re.search(r"When is the data collection for the research intended to begin and end\?.*?\n\s*(\d{2}/\d{4})\s*to\s*(\d{2}/\d{4})", part_5_text)
    if not dates:
        results["errors"].append("Data collection dates are missing or improperly formatted.")
    else:
        start_date_str, end_date_str = dates.group(1), dates.group(2)
        start_date = validate_date_format(start_date_str)
        end_date = validate_date_format(end_date_str)
        
        if not start_date or not end_date:
            results["errors"].append("Data collection dates must be in MM/YYYY format.")
        else:
            results["info"].append(f"Data collection dates: {start_date_str} to {end_date_str}.")
            delta = (end_date.year - start_date.year) * 12 + end_date.month - start_date.month
            if delta > 12:
                results["errors"].append("Data collection period exceeds one year.")
    
    involvement = re.search(r"Describe how subjects will be involved in detail.*?\n(.*?)(Will you be the one administering|$)", part_5_text, re.DOTALL)
    involvement_text = ""
    if not involvement or not involvement.group(1).strip():
        results["errors"].append("Participant involvement description is missing or empty.")
    else:
        involvement_text = involvement.group(1).strip()
        results["info"].append("Participant involvement description provided.")
        if "debriefing" in involvement_text.lower():
            required_forms.append({"form": "Debriefing Documents", "reason": "Required if debriefing is part of the process."})
    
    administration = re.search(r"Will you be the one administering.*?\n(.*?)(Will the participants experience any discomfort|$)", part_5_text, re.DOTALL)
    if not administration or not administration.group(1).strip():
        results["errors"].append("Data collection administration description is missing or empty.")
    else:
        results["info"].append("Data collection administration description provided.")
    
    discomfort = re.search(r"Will the participants experience any discomfort\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_5_text)
    discomfort_na = re.search(r"If “Yes”, please explain.*?\n.*?(N/A ☑|N/A ☐)", part_5_text, re.DOTALL)
    discomfort_explanation = re.search(r"If “Yes”, please explain.*?\n(.*?)(Will deception or false or misleading|$)", part_5_text, re.DOTALL)
    
    if not discomfort:
        results["errors"].append("Discomfort question not found or improperly formatted.")
    else:
        discomfort_answer = discomfort.group(1)
        if "☐" in discomfort_answer:
            results["warnings"].append("Discomfort checkbox is not marked (☐).")
        elif discomfort_answer == "Yes ☑":
            results["info"].append("Discomfort: Yes.")
            required_forms.append({"form": "Appendix B: Written Informed Consent Form", "reason": "Required for discomfort, with precautions."})
            if discomfort_na and discomfort_na.group(1) == "N/A ☑":
                results["errors"].append("Discomfort explanation cannot be N/A when discomfort is Yes.")
            elif not discomfort_explanation or not discomfort_explanation.group(1).strip():
                results["errors"].append("Discomfort explanation is missing or empty.")
            else:
                results["info"].append("Discomfort explanation provided.")
        else:
            results["info"].append("Discomfort: No.")
            if discomfort_explanation and discomfort_explanation.group(1).strip():
                results["warnings"].append("Discomfort explanation provided when discomfort is No.")
    
    deception = re.search(r"Will deception or false or misleading information be used.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_5_text, re.DOTALL)
    deception_na = re.search(r"If “Yes”, explain why deception is necessary.*?\n.*?(N/A ☑|N/A ☐)", part_5_text, re.DOTALL)
    deception_explanation = re.search(r"If “Yes”, explain why deception is necessary.*?\n(.*?)(Part 6:|$)", part_5_text, re.DOTALL)
    
    if not deception:
        results["errors"].append("Deception question not found or improperly formatted.")
    else:
        deception_answer = deception.group(1)
        if "☐" in deception_answer:
            results["warnings"].append("Deception checkbox is not marked (☐).")
        elif deception_answer == "Yes ☑":
            results["info"].append("Deception: Yes.")
            required_forms.extend([
                {"form": "Appendix B: Written Informed Consent Form", "reason": "Required for deception, with debriefing."},
                {"form": "Debriefing Documents", "reason": "Required for deception."}
            ])
            if deception_na and deception_na.group(1) == "N/A ☑":
                results["errors"].append("Deception explanation cannot be N/A when deception is Yes.")
            elif not deception_explanation or not deception_explanation.group(1).strip():
                results["errors"].append("Deception explanation is missing or empty.")
            else:
                results["info"].append("Deception explanation provided.")
        else:
            results["info"].append("Deception: No.")
            if deception_explanation and deception_explanation.group(1).strip():
                results["warnings"].append("Deception explanation provided when deception is No.")
    
    return results, required_forms, involvement_text

def validate_part_6(doc: docx.Document, required_forms: List[Dict]) -> Tuple[Dict[str, List[str]], List[Dict], str, str, str]:
    """Validates Part 6: Data Management Plan."""
    results = {"errors": [], "warnings": [], "info": []}
    part_6_text = ""
    in_part_6 = False
    
    for para in doc.paragraphs:
        if "Part 6: Data Management Plan" in para.text:
            in_part_6 = True
        if in_part_6:
            part_6_text += para.text + "\n"
        if "Part 7: Risk/Benefit Analysis" in para.text:
            break
    
    if not part_6_text:
        results["errors"].append("Part 6: Data Management Plan section not found.")
        return results, required_forms, "", "", ""
    
    electronic_survey = re.search(r"Are you conducting a survey using any electronic media\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_6_text)
    if not electronic_survey:
        results["errors"].append("Electronic survey question not found or improperly formatted.")
        return results, required_forms, "", "", ""
    
    survey_answer = electronic_survey.group(1)
    if "☐" in survey_answer:
        results["warnings"].append("Electronic survey checkbox is not marked (☐).")
        return results, required_forms, "", "", ""
    
    results["info"].append(f"Electronic survey: {survey_answer}.")
    if survey_answer == "Yes ☑":
        required_forms.append({"form": "Appendix C: Informed Consent Form for Internet Surveys", "reason": "Required for internet surveys."})
        
        name_privacy = re.search(r"Will you assure that the participant will only see his/her name\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_6_text)
        read_receipt = re.search(r"Will you have the “read receipt” function turned off\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_6_text)
        email_explanation = re.search(r"If you answered “No” to these questions, please explain.*?\n(.*?)(If your survey contains questions|$)", part_6_text, re.DOTALL)
        
        if not name_privacy:
            results["errors"].append("Name privacy question not found or improperly formatted.")
        else:
            results["info"].append(f"Name privacy: {name_privacy.group(1)}.")
        
        if not read_receipt:
            results["errors"].append("Read receipt question not found or improperly formatted.")
        else:
            results["info"].append(f"Read receipt: {read_receipt.group(1)}.")
        
        if (name_privacy and name_privacy.group(1) == "No ☑") or (read_receipt and read_receipt.group(1) == "No ☑"):
            if not email_explanation or not email_explanation.group(1).strip():
                results["errors"].append("Explanation for 'No' in email invitation questions is missing or empty.")
            else:
                results["info"].append("Explanation for 'No' in email invitation provided.")
        
        dropdown = re.search(r"Do they have the option to choose “No response”.*?(Yes ☑|No ☑|No dropdown menu ☑|Yes ☐|No ☐|No dropdown menu ☐)", part_6_text)
        if not dropdown:
            results["errors"].append("Dropdown menu question not found or improperly formatted.")
        else:
            results["info"].append(f"Dropdown menu: {dropdown.group(1)}.")
        
        transmission = re.search(r"How will data be transmitted\?.*?\n(.*?)(What is the URL\?|$)", part_6_text, re.DOTALL)
        if not transmission or not transmission.group(1).strip():
            results["errors"].append("Data transmission description is missing or empty.")
        else:
            results["info"].append("Data transmission description provided.")
        
        url = re.search(r"What is the URL\?.*?\n\s*([^\n]*)", part_6_text)
        if not url or not url.group(1).strip():
            results["errors"].append("URL is missing or empty for electronic survey.")
        else:
            results["info"].append(f"Survey URL: {url.group(1).strip()}.")
    
    else:
        for field, pattern in [
            ("Name privacy", r"Will you assure that the participant will only see his/her name\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
            ("Read receipt", r"Will you have the “read receipt” function turned off\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)"),
            ("Dropdown menu", r"Do they have the option to choose “No response”.*?(Yes ☑|No ☑|No dropdown menu ☑|Yes ☐|No ☐|No dropdown menu ☐)"),
            ("Data transmission", r"How will data be transmitted\?.*?\n(.*?)(What is the URL\?|$)"),
            ("URL", r"What is the URL\?.*?\n\s*([^\n]*)")
        ]:
            match = re.search(pattern, part_6_text, re.DOTALL)
            if match and ((field in ["Name privacy", "Read receipt", "Dropdown menu"] and match.group(1) not in ["Yes ☐", "No ☐", "No dropdown menu ☐"]) or
                         (field in ["Data transmission", "URL"] and match.group(1).strip())):
                results["errors"].append(f"{field} should be unanswered or empty when electronic survey is No.")
    
    storage = re.search(r"Where will data be stored\?.*?\n\s*([^\n]*)", part_6_text)
    storage_text = ""
    if not storage or not storage.group(1).strip():
        results["errors"].append("Data storage description is missing or empty.")
    else:
        storage_text = storage.group(1).strip()
        results["info"].append(f"Data storage: {storage_text}.")
    
    maintenance = re.search(r"How will data be maintained\?.*?\n(.*?)(Will data be shared\?|$)", part_6_text, re.DOTALL)
    maintenance_text = ""
    if not maintenance or not maintenance.group(1).strip():
        results["errors"].append("Data maintenance description is missing or empty.")
    else:
        maintenance_text = maintenance.group(1).strip()
        results["info"].append(f"Data maintenance: {maintenance_text}.")
        if "identifiable" in maintenance_text.lower():
            required_forms.append({"form": "Appendix L: Confidentiality Agreement Form", "reason": "Recommended for identifiable data."})
    
    sharing = re.search(r"Will data be shared\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_6_text)
    sharing_details = re.search(r"How\? With whom\? Will subjects be re-identifiable\? Why or why not\?.*?\n(.*?)(Describe the data security plan|$)", part_6_text, re.DOTALL)
    sharing_text = ""
    
    if not sharing:
        results["errors"].append("Data sharing question not found or improperly formatted.")
    else:
        sharing_answer = sharing.group(1)
        if "☐" in sharing_answer:
            results["warnings"].append("Data sharing checkbox is not marked (☐).")
        else:
            results["info"].append(f"Data sharing: {sharing_answer}.")
            if not sharing_details or not sharing_details.group(1).strip():
                results["errors"].append("Data sharing details are missing or empty.")
            else:
                sharing_text = sharing_details.group(1).strip()
                results["info"].append(f"Data sharing details: {sharing_text}.")
                if sharing_answer == "Yes ☑" and "identifiable" in sharing_text.lower():
                    required_forms.append({"form": "Appendix L: Confidentiality Agreement Form", "reason": "Required for sharing identifiable data."})
    
    security = re.search(r"Describe the data security plan.*?\n(.*?)(Part 7:|$)", part_6_text, re.DOTALL)
    if not security or not security.group(1).strip():
        results["errors"].append("Data security plan description is missing or empty.")
    else:
        results["info"].append(f"Data security plan: {security.group(1).strip()}.")
    
    return results, required_forms, maintenance_text, sharing_text, storage_text

def validate_part_7(doc: docx.Document, required_forms: List[Dict]) -> Tuple[Dict[str, List[str]], List[Dict]]:
    """Validates Part 7: Risk/Benefit Analysis."""
    results = {"errors": [], "warnings": [], "info": []}
    part_7_text = ""
    in_part_7 = False
    
    for para in doc.paragraphs:
        if "Part 7: Risk/Benefit Analysis" in para.text:
            in_part_7 = True
        if in_part_7:
            part_7_text += para.text + "\n"
        if "Part 8: Confidentiality/Anonymity" in para.text:
            break
    
    if not part_7_text:
        results["errors"].append("Part 7: Risk/Benefit Analysis section not found.")
        return results, required_forms
    
    minimal_risk = re.search(r"Do you believe those risks will be no greater than minimal\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_7_text)
    minimal_risk_explanation = re.search(r"Explain why:.*?\n(.*?)(Describe all risks|$)", part_7_text, re.DOTALL)
    
    if not minimal_risk:
        results["errors"].append("Minimal risk question not found or improperly formatted.")
    else:
        minimal_risk_answer = minimal_risk.group(1)
        if "☐" in minimal_risk_answer:
            results["warnings"].append("Minimal risk checkbox is not marked (☐).")
        else:
            results["info"].append(f"Minimal risk: {minimal_risk_answer}.")
        
        if not minimal_risk_explanation or not minimal_risk_explanation.group(1).strip():
            results["errors"].append("Minimal risk explanation is missing or empty.")
        else:
            results["info"].append("Minimal risk explanation provided.")
    
    risks = re.search(r"Describe all risks.*?\n(.*?)(If risks are greater than minimal|$)", part_7_text, re.DOTALL)
    if not risks or not risks.group(1).strip():
        results["errors"].append("Risks description is missing or empty.")
    else:
        risks_text = risks.group(1).strip()
        if risks_text.lower() in ["not applicable", "no risk"]:
            results["errors"].append("Risks description cannot be 'Not Applicable' or 'No risk'.")
        else:
            results["info"].append(f"Risks description: {risks_text}.")
    
    if minimal_risk and minimal_risk.group(1) == "No ☑":
        required_forms.append({"form": "Appendix B: Written Informed Consent Form", "reason": "Required for greater than minimal risk."})
        risk_fields = [
            ("Why risks are essential", r"Explain why these risks are essential.*?\n(.*?)(What have you done to minimize risks|$)", part_7_text),
            ("Minimize risks", r"What have you done to minimize risks.*?\n(.*?)(What protections have you put in place|$)", part_7_text),
            ("Protections", r"What protections have you put in place.*?\n(.*?)(What procedures have you established|$)", part_7_text),
            ("Adverse events reporting", r"What procedures have you established for reporting adverse events.*?\n(.*?)(Will the participants directly|$)", part_7_text)
        ]
        
        for field_name, pattern, text in risk_fields:
            match = re.search(pattern, text, re.DOTALL)
            if not match or not match.group(1).strip():
                results["errors"].append(f"{field_name} description is missing or empty.")
            else:
                results["info"].append(f"{field_name} description provided.")
    
    participant_benefits = re.search(r"Will the participants directly or indirectly benefit.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_7_text, re.DOTALL)
    benefits_explanation = re.search(r"Please explain:.*?\n(.*?)(What are the anticipated benefits to society|$)", part_7_text, re.DOTALL)
    
    if not participant_benefits:
        results["errors"].append("Participant benefits question not found or improperly formatted.")
    else:
        benefits_answer = participant_benefits.group(1)
        if "☐" in benefits_answer:
            results["warnings"].append("Participant benefits checkbox is not marked (☐).")
        else:
            results["info"].append(f"Participant benefits: {benefits_answer}.")
            if not benefits_explanation or not benefits_explanation.group(1).strip():
                results["errors"].append("Participant benefits explanation is missing or empty.")
            else:
                results["info"].append("Participant benefits explanation provided.")
                if benefits_answer == "Yes ☑":
                    required_forms.append({
                        "form": "Appendix B: Written Informed Consent Form",
                        "reason": "Required to detail participant benefits."
                    })
    
    societal_benefits = re.search(r"What are the anticipated benefits to society.*?\n(.*?)(Will incentives be offered|$)", part_7_text, re.DOTALL)
    if not societal_benefits or not societal_benefits.group(1).strip():
        results["errors"].append("Societal benefits description is missing or empty.")
    else:
        results["info"].append("Societal benefits description provided.")
    
    incentives = re.search(r"Will incentives be offered.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_7_text, re.DOTALL)
    incentives_details = re.search(r"If “Yes”, please describe.*?\n(.*?)(Part 8:|$)", part_7_text, re.DOTALL)
    
    if not incentives:
        results["errors"].append("Incentives question not found or improperly formatted.")
    else:
        incentives_answer = incentives.group(1)
        if "☐" in incentives_answer:
            results["warnings"].append("Incentives checkbox is not marked (☐).")
        elif incentives_answer == "Yes ☑":
            results["info"].append("Incentives: Yes.")
            required_forms.append({
                "form": "Appendix B: Written Informed Consent Form",
                "reason": "Required for incentives, with details."
            })
            if not incentives_details or not incentives_details.group(1).strip():
                results["errors"].append("Incentives description is missing or empty.")
            else:
                results["info"].append("Incentives description provided.")
        else:
            results["info"].append("Incentives: No.")
            if incentives_details and incentives_details.group(1).strip():
                results["warnings"].append("Incentives description provided when incentives is No.")
    
    return results, required_forms

def validate_part_8(doc: docx.Document, required_forms: List[Dict], methodology_text: str, involvement_text: str, maintenance_text: str, sharing_text: str, storage_text: str) -> Tuple[Dict[str, List[str]], List[Dict]]:
    """Validates Part 8: Confidentiality/Anonymity with consistency checks."""
    results = {"errors": [], "warnings": [], "info": []}
    part_8_text = ""
    in_part_8 = False
    
    for para in doc.paragraphs:
        if "Part 8: Confidentiality/Anonymity" in para.text:
            in_part_8 = True
        if in_part_8:
            part_8_text += para.text + "\n"
        if "Part 10: Project Funding" in para.text:
            break
    
    if not part_8_text:
        results["errors"].append("Part 8: Confidentiality/Anonymity section not found.")
        return results, required_forms
    
    recordings = re.search(r"Will you be video recording.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_8_text, re.DOTALL)
    if not recordings:
        results["errors"].append("Recordings question not found or improperly formatted.")
    else:
        recordings_answer = recordings.group(1)
        if "☐" in recordings_answer:
            results["warnings"].append("Recordings checkbox is not marked (☐).")
        else:
            results["info"].append(f"Video/Photograph/Audio Recordings: {recordings_answer}.")
            if recordings_answer == "Yes ☑":
                required_forms.append({
                    "form": "Appendix B: Written Informed Consent Form",
                    "reason": "Required for recordings."
                })
            
            # Consistency check with Parts 3 and 5
            recording_terms = ["video", "audio", "photograph", "recording", "interview via video"]
            if any(term in methodology_text.lower() for term in recording_terms) or any(term in involvement_text.lower() for term in recording_terms):
                if recordings_answer != "Yes ☑":
                    results["errors"].append("Part 8.1 should be 'Yes' as Parts 3 or 5 mention video/audio/photograph.")
            elif recordings_answer == "Yes ☑":
                results["warnings"].append("Part 8.1 is 'Yes' but no video/audio/photograph mentioned in Parts 3 or 5.")
    
    consent_recordings = re.search(r"Will you be obtaining signed consent forms.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_8_text, re.DOTALL)
    if not consent_recordings:
        results["errors"].append("Consent for recordings question not found or improperly formatted.")
    else:
        consent_answer = consent_recordings.group(1)
        if "☐" in consent_answer:
            results["warnings"].append("Consent for recordings checkbox is not marked (☐).")
        else:
            results["info"].append(f"Consent for recordings: {consent_answer}.")
            if recordings_answer == "Yes ☑" and consent_answer != "Yes ☑":
                results["errors"].append("Consent for recordings must be 'Yes' when recordings is 'Yes'.")
            if recordings_answer == "No ☑" and consent_answer == "Yes ☑":
                results["errors"].append("Consent for recordings should be 'No' or unanswered when recordings is 'No'.")
    
    identifiability = re.search(r"Will the data be identifiable.*?\n.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_8_text, re.DOTALL)
    identifiability_explanation = re.search(r"If “Yes”, please explain.*?\n(.*?)(Describe procedures to create/preserve anonymity|$)", part_8_text, re.DOTALL)
    
    if not identifiability:
        results["errors"].append("Identifiability question not found or improperly formatted.")
    else:
        identifiability_answer = identifiability.group(1)
        if "☐" in identifiability_answer:
            results["warnings"].append("Identifiability checkbox is not marked (☐).")
        else:
            results["info"].append(f"Identifiability: {identifiability_answer}.")
            if identifiability_answer == "Yes ☑":
                required_forms.append({
                    "form": "Appendix L: Confidentiality Agreement Form",
                    "reason": "Required for identifiable data."
                })
                if not identifiability_explanation or not identifiability_explanation.group(1).strip():
                    results["errors"].append("Identifiability explanation is missing or empty when identifiability is Yes.")
                else:
                    results["info"].append("Identifiability explanation provided.")
                
                # Consistency check with Part 6
                if any("identifiable" in text.lower() for text in [maintenance_text, sharing_text, storage_text]):
                    results["info"].append("Identifiability in Part 8 is consistent with Part 6.")
                else:
                    results["warnings"].append("Part 8.3 is 'Yes' but Part 6 does not mention identifiable data.")
            else:
                if any("identifiable" in text.lower() for text in [maintenance_text, sharing_text, storage_text]):
                    results["errors"].append("Part 8.3 should be 'Yes' as Part 6 mentions identifiable data.")
    
    anonymity_na = re.search(r"Describe procedures to create/preserve anonymity.*?\n.*?(N/A ☑|N/A ☐)", part_8_text, re.DOTALL)
    anonymity_procedures = re.search(r"Describe procedures to create/preserve anonymity.*?\n(.*?)(Describe procedures to preserve confidentiality|$)", part_8_text, re.DOTALL)
    
    if identifiability_answer == "No ☑":
        if anonymity_na and anonymity_na.group(1) == "N/A ☑":
            results["errors"].append("Anonymity procedures cannot be N/A when identifiability is No.")
        elif not anonymity_procedures or not anonymity_procedures.group(1).strip():
            results["errors"].append("Anonymity procedures description is missing or empty when identifiability is No.")
        else:
            results["info"].append("Anonymity procedures description provided.")
    else:
        if anonymity_na and anonymity_na.group(1) != "N/A ☑":
            results["errors"].append("Anonymity procedures should be N/A when identifiability is Yes.")
        if anonymity_procedures and anonymity_procedures.group(1).strip():
            results["errors"].append("Anonymity procedures should be empty when identifiability is Yes.")
    
    confidentiality_fields = [
        ("During data collection", r"During data collection.*?\n(.*?)(While results are analyzed|$)", part_8_text),
        ("While results are analyzed", r"While results are analyzed.*?\n(.*?)(In publication/reporting|$)", part_8_text),
        ("In publication/reporting", r"In publication/reporting.*?\n(.*?)(In storage after research completion|$)", part_8_text),
        ("In storage after research completion", r"In storage after research completion.*?\n(.*?)(Part 10:|$)", part_8_text)
    ]
    
    if identifiability_answer == "Yes ☑":
        for field_name, pattern, text in confidentiality_fields:
            match = re.search(pattern, text, re.DOTALL)
            if not match or not match.group(1).strip():
                results["errors"].append(f"Confidentiality procedures for '{field_name}' are missing or empty.")
            else:
                results["info"].append(f"Confidentiality procedures for '{field_name}' provided.")
    else:
        for field_name, pattern, text in confidentiality_fields:
            match = re.search(pattern, text, re.DOTALL)
            if match and match.group(1).strip():
                results["warnings"].append(f"Confidentiality procedures for '{field_name}' should be empty when identifiability is No.")
    
    return results, required_forms

def validate_part_10(doc: docx.Document, required_forms: List[Dict]) -> Tuple[Dict[str, List[str]], List[Dict]]:
    """Validates Part 10: Project Funding."""
    results = {"errors": [], "warnings": [], "info": []}
    part_10_text = ""
    in_part_10 = False
    
    for para in doc.paragraphs:
        if "Part 10: Project Funding" in para.text:
            in_part_10 = True
        if in_part_10:
            part_10_text += para.text + "\n"
        if "Part 11: Protocol for naming of documents" in para.text:
            break
    
    if not part_10_text:
        results["errors"].append("Part 10: Project Funding section not found.")
        return results, required_forms
    
    funding = re.search(r"Is this project being supported by any funding sources\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_10_text)
    if not funding:
        results["errors"].append("Funding question not found or improperly formatted.")
    else:
        funding_answer = funding.group(1)
        if "☐" in funding_answer:
            results["warnings"].append("Funding checkbox is not marked (☐).")
        else:
            results["info"].append(f"Project funding: {funding_answer}.")
            if funding_answer == "Yes ☑":
                source = re.search(r"If yes, please specify the funding source\(s\):.*?\n(.*?)(Is the funding external|$)", part_10_text, re.DOTALL)
                external = re.search(r"Is the funding external to Nazarbayev University\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_10_text)
                
                if not source or not source.group(1).strip():
                    results["errors"].append("Funding source description is missing or empty.")
                else:
                    results["info"].append(f"Funding source: {source.group(1).strip()}.")
                
                if not external:
                    results["errors"].append("External funding question not found or improperly formatted.")
                else:
                    external_answer = external.group(1)
                    if "☐" in external_answer:
                        results["warnings"].append("External funding checkbox is not marked (☐).")
                    else:
                        results["info"].append(f"External funding: {external_answer}.")
                        if external_answer == "Yes ☑":
                            required_forms.append({
                                "form": "Appendix K: Funding Source Form",
                                "reason": "Required for external funding."
                            })
            else:
                source = re.search(r"If yes, please specify the funding source\(s\):.*?\n(.*?)(Is the funding external|$)", part_10_text, re.DOTALL)
                external = re.search(r"Is the funding external to Nazarbayev University\?.*?(Yes ☑|No ☑|Yes ☐|No ☐)", part_10_text)
                if source and source.group(1).strip():
                    results["errors"].append("Funding source should be empty when funding is No.")
                if external and external.group(1) not in ["Yes ☐", "No ☐"]:
                    results["errors"].append("External funding question should be unanswered when funding is No.")
    
    return results, required_forms

def validate_part_11_and_checklist(doc: docx.Document, required_forms: List[Dict], pi_surname: str, file_names: Optional[List[str]] = None) -> Tuple[Dict[str, List[str]], List[Dict]]:
    """Validates Part 11: Protocol for naming of documents and Checklist."""
    results = {"errors": [], "warnings": [], "info": []}
    part_11_text = ""
    in_part_11 = False
    
    for para in doc.paragraphs:
        if "Part 11: Protocol for naming of documents" in para.text:
            in_part_11 = True
        if in_part_11:
            part_11_text += para.text + "\n"
    
    if not part_11_text:
        results["errors"].append("Part 11: Protocol for naming of documents section not found.")
        return results, required_forms
    
    # Validate document naming protocol
    if not pi_surname:
        results["errors"].append("PI surname is missing; cannot validate file naming protocol.")
    else:
        if not file_names:
            results["warnings"].append("No file names provided for naming protocol validation.")
        else:
            application_found = False
            for file_name in file_names:
                is_valid, message = validate_file_name(file_name, pi_surname)
                if is_valid:
                    results["info"].append(message)
                    if "Application form naming is valid" in message:
                        application_found = True
                else:
                    results["errors"].append(message)
            
            if not application_found:
                results["errors"].append("Main application file (Surname_IREC Application_MMDDYYYY) not found in provided file names.")
    
    # Validate checklist
    checklist_start = re.search(r"CHECKLIST\s*Please indicate which forms.*?\n", part_11_text, re.DOTALL)
    if not checklist_start:
        results["errors"].append("Checklist section not found in Part 11.")
        return results, required_forms
    
    checklist_text = part_11_text[checklist_start.end():]
    
    # Extract all checklist items
    checklist_items = re.findall(r"([^\n]+?)\s*(☑|☐)\s*(?:\n|$)", checklist_text, re.DOTALL)
    if not checklist_items:
        results["errors"].append("No checklist items found in Part 11.")
        return results, required_forms
    
    # Create a set of checked forms
    checked_forms = {item.strip(): status for item, status in checklist_items}
    
    # Validate required forms
    for required_form in required_forms:
        form_name = required_form["form"].strip()
        if form_name not in checked_forms:
            results["errors"].append(f"Required form '{form_name}' not listed in checklist.")
        elif checked_forms[form_name] != "☑":
            results["errors"].append(f"Required form '{form_name}' is listed but not checked (☐).")
        else:
            results["info"].append(f"Required form '{form_name}' is correctly checked (☑).")
    
    # Check for unnecessary forms
    required_form_names = {form["form"].strip() for form in required_forms}
    for form_name, status in checked_forms.items():
        if form_name not in required_form_names and status == "☑":
            results["warnings"].append(f"Form '{form_name}' is checked but not required.")
    
    return results, required_forms

def validate_irec_application(doc: docx.Document, file_names: Optional[List[str]] = None) -> Dict[str, any]:
    """
    Validates the entire NU IREC application by calling validation functions for each part.
    Returns a consolidated report with submission ID, timestamp, and validation results.
    """
    results = {
        "submission_id": str(uuid.uuid4()),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "parts": {},
        "summary": {"errors": 0, "warnings": 0, "info": 0}
    }
    
    required_forms = [
        {"form": "Appendix A: IREC Application Form", "reason": "Required for all submissions."},
        {"form": "CITI Training Certificates", "reason": "Required for all team members."}
    ]
    
    # Validate Part 0
    part_0_results, exemption_claimed = validate_part_0(doc)
    results["parts"]["Part 0"] = part_0_results
    
    # Validate Part 1
    part_1_results = validate_part_1(doc, exemption_claimed)
    results["parts"]["Part 1"] = part_1_results
    
    # Validate Part 2 and extract PI surname
    part_2_results, pi_surname = validate_part_2(doc)
    results["parts"]["Part 2"] = part_2_results
    
    # Validate Part 3
    part_3_results, required_forms, methodology_text = validate_part_3(doc)
    results["parts"]["Part 3"] = part_3_results
    
    # Validate Part 4
    part_4_results, required_forms = validate_part_4(doc, required_forms)
    results["parts"]["Part 4"] = part_4_results
    
    # Validate Part 5
    part_5_results, required_forms, involvement_text = validate_part_5(doc, required_forms)
    results["parts"]["Part 5"] = part_5_results
    
    # Validate Part 6
    part_6_results, required_forms, maintenance_text, sharing_text, storage_text = validate_part_6(doc, required_forms)
    results["parts"]["Part 6"] = part_6_results
    
    # Validate Part 7
    part_7_results, required_forms = validate_part_7(doc, required_forms)
    results["parts"]["Part 7"] = part_7_results
    
    # Validate Part 8
    part_8_results, required_forms = validate_part_8(doc, required_forms, methodology_text, involvement_text, maintenance_text, sharing_text, storage_text)
    results["parts"]["Part 8"] = part_8_results
    
    # Validate Part 10
    part_10_results, required_forms = validate_part_10(doc, required_forms)
    results["parts"]["Part 10"] = part_10_results
    
    # Validate Part 11 and Checklist
    part_11_results, required_forms = validate_part_11_and_checklist(doc, required_forms, pi_surname, file_names)
    results["parts"]["Part 11"] = part_11_results
    
    # Summarize results
    for part, part_results in results["parts"].items():
        results["summary"]["errors"] += len(part_results["errors"])
        results["summary"]["warnings"] += len(part_results["warnings"])
        results["summary"]["info"] += len(part_results["info"])
    
    # Add list of required forms to results
    results["required_forms"] = required_forms
    
    return results

if __name__ == "__main__":
    # Example usage
    try:
        doc = docx.Document("irec_application.docx")
        file_names = [
            "Smith_IREC Application_01012025.docx",
            "Smith_InterviewQuestions-Eng_01012025.pdf",
            "Smith_ConsentForm-Eng_01012025.docx",
            "Smith_CITI_01012023.pdf"
        ]
        validation_results = validate_irec_application(doc, file_names)
        
        print(f"Submission ID: {validation_results['submission_id']}")
        print(f"Timestamp: {validation_results['timestamp']}")
        print("\nValidation Summary:")
        print(f"Total Errors: {validation_results['summary']['errors']}")
        print(f"Total Warnings: {validation_results['summary']['warnings']}")
        print(f"Total Info Messages: {validation_results['summary']['info']}")
        
        for part, results in validation_results["parts"].items():
            print(f"\n{part}:")
            for category, messages in results.items():
                if messages:
                    print(f"  {category.capitalize()}:")
                    for msg in messages:
                        print(f"    - {msg}")
        
        print("\nRequired Forms:")
        for form in validation_results["required_forms"]:
            print(f"  - {form['form']}: {form['reason']}")
        
    except FileNotFoundError:
        print("Error: 'irec_application.docx' not found.")
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        
